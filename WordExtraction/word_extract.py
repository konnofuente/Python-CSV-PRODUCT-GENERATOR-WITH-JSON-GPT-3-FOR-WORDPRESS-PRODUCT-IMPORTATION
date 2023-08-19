# from docx import Document

# # Load the Word document
# doc = Document('H:\My Drive\Client\La rosee\Web developement\code\WordExtraction\BARQUETTES ASSIETTES COUVERTS.docx')

# # Extract the text from each paragraph in the document
# text = []
# for paragraph in doc.paragraphs:
#     text.append(paragraph.text)

# # Join the extracted text into a single string
# document_text = '\n'.join(text)

# # Print the extracted text
# print(document_text)



# import csv
# from docx import Document

# # Open the Word document
# doc = Document('H:\My Drive\Client\La rosee\Web developement\code\WordExtraction\BARQUETTES ASSIETTES COUVERTS.docx')

# # Create a list to store the extracted data
# data = []

# # Iterate through each paragraph in the document
# for paragraph in doc.paragraphs:
#     # Check if the paragraph contains product details
#     if 'product code' in paragraph.text.lower():
#         # Extract the product name, code, and price
#         product_name = paragraph.text.split(':')[0].strip()
#         product_code = paragraph.text.split(':')[1].split(')')[0].strip().replace('product code', '')
#         product_price = paragraph.text.split('XAF')[1].strip()

#         # Append the extracted data to the list
#         data.append([product_name, product_code, product_price])

# # Write the extracted data to a CSV file
# with open('output12.csv', 'w', newline='') as csvfile:
#     writer = csv.writer(csvfile)
#     writer.writerow(['Product Name', 'Product Code', 'Product Price'])
#     writer.writerows(data)



# import csv
# from docx import Document
# import re

# # Open the Word document
# doc = Document('H:\My Drive\Client\La rosee\Web developement\code\WordExtraction\BARQUETTES ASSIETTES COUVERTS.docx')

# # Create a list to store the extracted data
# data = []

# # Iterate through each paragraph in the document
# for paragraph in doc.paragraphs:
#     # Check if the paragraph contains product details
#     if 'product code' in paragraph.text.lower():
#         # Extract the product name, code, and price
#         product_name_match = re.search(r'(.*?)\(', paragraph.text)
#         if product_name_match:
#             product_name = product_name_match.group(1).strip()
#         else:
#             product_name = ''
#         product_code = paragraph.text.split(':')[1].split(')')[0].strip().replace('product code', '')
#         product_price = paragraph.text.split('XAF')[1].strip()

#         # Extract the detail from the paragraph before "FICHE TECHNIQUE"
#         detail = re.search(r'(.*?)FICHE TECHNIQUE', paragraph.text, re.DOTALL)
#         if detail:
#             detail = detail.group(1).strip()
#         else:
#             detail = ''

#         # Append the extracted data to the list
#         data.append([product_name, product_code, product_price, 'XAF', detail])

# # Write the extracted data to a CSV file
# with open('output2.csv', 'w', newline='') as csvfile:
#     writer = csv.writer(csvfile)
#     writer.writerow(['Product Name', 'Product Code', 'Product Price', 'Currency', 'Detail'])
#     writer.writerows(data)



import csv
from docx import Document
import re

# Open the Word document
doc = Document('H:\My Drive\Client\La rosee\Web developement\code\WordExtraction\BARQUETTES ASSIETTES COUVERTS.docx')

# Create a list to store the extracted data
data = []

# Iterate through each paragraph in the document
for paragraph in doc.paragraphs:
    # Check if the paragraph contains product details
    if 'product code' in paragraph.text.lower():
        # Extract the product name, code, and price
        product_name_match = re.search(r'(.*?)\(', paragraph.text)
        if product_name_match:
            product_name = product_name_match.group(1).strip()
        else:
            product_name = ''
        product_code = paragraph.text.split(':')[1].split(')')[0].strip().replace('product code', '')
        product_price = paragraph.text.split('XAF')[1].strip()

        # Extract the detail from the paragraph between the price and "FICHE TECHNIQUE"
        detail = re.search(r'{price}(.*?){fiche_technique}'.format(price=product_price, fiche_technique='FICHE TECHNIQUE'), paragraph.text, re.DOTALL)
        if detail:
            detail = detail.group(1).strip()
        else:
            detail = ''

        # Append the extracted data to the list
        data.append([product_name, product_code, product_price, 'XAF', detail])

# Write the extracted data to a CSV file
with open('audray.csv', 'w', newline='') as csvfile:
    writer = csv.writer(csvfile)
    writer.writerow(['Product Name', 'Product Code', 'Product Price', 'Currency', 'Detail'])
    writer.writerows(data)


# import spacy
# import json
# from docx import Document

# # Load the spaCy NER model
# nlp = spacy.load("en_core_web_sm")

# # Read the document file
# document = Document('document.docx')
# document_text = ' '.join([paragraph.text for paragraph in document.paragraphs])

# # Process the document with the NER model
# doc = nlp(document_text)

# # Initialize empty lists for product details
# product_details = []

# # Iterate over the entities in the document
# for ent in doc.ents:
#     # Check if the entity is a product name
#     if ent.label_ == "PRODUCT_NAME":
#         # Find the next token after the product name
#         next_token = ent.end + 1
#         # Initialize empty dictionary for product details
#         details = {
#             "Product Name": ent.text,
#             "Product Code": "",
#             "Product Price": "",
#             "Currency": "",
#             "Detail": ""
#         }
#         # Loop through the subsequent tokens to find the required details
#         while next_token < len(doc):
#             token = doc[next_token]
#             # Check if the token is a product code
#             if token.ent_type_ == "PRODUCT_CODE":
#                 details["Product Code"] = token.text
#             # Check if the token is a price
#             elif token.ent_type_ == "PRODUCT_PRICE":
#                 details["Product Price"] = token.text
#                 details["Currency"] = token.ent_type_
#             # Check if the token is a detail
#             elif token.ent_type_ == "DETAIL":
#                 details["Detail"] = token.text
#                 break
#             next_token += 1
#         product_details.append(details)

# # Write the data to a JSON file
# with open('output20.json', 'w') as file:
#     json.dump(product_details, file)
